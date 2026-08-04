from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\lenovo\Desktop\OCR total")
OUTPUT = ROOT / "OCR项目环境配置与运行指南.docx"


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.3)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F1F5F8")
    p._p.get_or_add_pPr().append(shading)
    for index, line in enumerate(text.splitlines()):
        if index:
            p.add_run().add_break()
        set_run_font(p.add_run(line), "Consolas", 8.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_font(run, size=15 if level == 1 else 11.5, bold=True, color=(23, 78, 122))
    return p


def add_text(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(item))


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.7)
section.right_margin = Cm(1.7)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(title.add_run("OCR 项目环境配置与运行指南"), size=22, bold=True, color=(18, 58, 99))
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(subtitle.add_run(r"适用目录：C:\Users\lenovo\Desktop\OCR total　｜　Windows 11"), size=10, color=(90, 100, 110))

add_heading(doc, "1. 项目结构与端口")
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
headers = ["组件", "目录", "技术与版本", "端口"]
for cell, value in zip(table.rows[0].cells, headers):
    set_run_font(cell.paragraphs[0].add_run(value), bold=True)
rows = [
    ("Python OCR 服务", "ocr-python-service", "Python 3.10/3.11、Flask、PaddleOCR", "5001"),
    ("Java 业务后端", "ocr-business-server", "JDK 17、Maven 3.9+、Spring Boot", "8080"),
    ("Web 前端", "ocr-web", "Node.js LTS、npm、Vue 3、Vite", "5173"),
]
for values in rows:
    cells = table.add_row().cells
    for cell, value in zip(cells, values):
        set_run_font(cell.paragraphs[0].add_run(value), size=9.5)
add_text(doc, "正确启动顺序：Python OCR 服务 → Java 后端 → Vue 前端。三个服务分别占用一个终端窗口，运行期间不要关闭。", "正确启动顺序：")

add_heading(doc, "2. Java 17 配置")
add_text(doc, "JDK 安装目录：")
add_code(doc, r"C:\Users\lenovo\AppData\Local\Programs\Eclipse Adoptium\jdk-17.0.20.8-hotspot")
add_text(doc, "用户环境变量：")
add_code(doc, "JAVA_HOME=C:\\Users\\lenovo\\AppData\\Local\\Programs\\Eclipse Adoptium\\jdk-17.0.20.8-hotspot\nPath 新增：%JAVA_HOME%\\bin")
add_text(doc, "验证：")
add_code(doc, "where java\njava -version")
add_text(doc, "应优先找到 Adoptium 路径并显示 Java 17。若仍找到 Java 8，检查系统 Path 中是否残留 Oracle 的 java8path 或 javapath。")

add_heading(doc, "3. Maven 配置")
add_code(doc, "MAVEN_HOME=C:\\Users\\lenovo\\Desktop\\apache-maven-3.9.16\nPath 新增：C:\\Users\\lenovo\\Desktop\\apache-maven-3.9.16\\bin")
add_code(doc, "echo %MAVEN_HOME%\nwhere mvn\nmvn -version")
add_text(doc, "结果应包含 Maven 3.9.16 和 Java 17。若在 VS Code 内修改环境变量，需要彻底关闭所有 Code.exe 后重新打开。")

doc.add_section(WD_SECTION.NEW_PAGE)
add_heading(doc, "4. Python OCR 环境")
add_text(doc, "推荐 Python 3.10 或 3.11（64 位），不要使用 Python 3.13。当前虚拟环境名为 jrrg。")
add_heading(doc, "激活环境并安装依赖", 2)
add_code(doc, 'conda activate jrrg\npython --version\ncd /d "C:\\Users\\lenovo\\Desktop\\OCR total\\ocr-python-service"\npython -m pip install --upgrade pip\npython -m pip install -r requirements.txt')
add_text(doc, "不在 Python 服务目录时，可以使用完整路径：")
add_code(doc, 'python -m pip install -r "C:\\Users\\lenovo\\Desktop\\OCR total\\ocr-python-service\\requirements.txt"')
add_heading(doc, "创建 .env 配置", 2)
add_code(doc, 'cd /d "C:\\Users\\lenovo\\Desktop\\OCR total\\ocr-python-service"\ncopy .env.example .env')
add_text(doc, "没有安装 Tesseract 时：")
add_code(doc, "TESSERACT_CMD=\nVERIFY_WITH_TESSERACT=false")
add_text(doc, "已安装 Tesseract 及中英文语言包时：")
add_code(doc, "TESSERACT_CMD=C:\\Program Files\\Tesseract-OCR\\tesseract.exe\nTESSERACT_LANG=chi_sim+eng\nVERIFY_WITH_TESSERACT=true")
add_heading(doc, "Windows CPU 的 oneDNN 兼容处理", 2)
add_text(doc, "出现 ConvertPirAttribute2RuntimeAttribute not support 时，在 .env 增加：")
add_code(doc, "FLAGS_use_mkldnn=0")
add_text(doc, "若仍报错，在 ocr_service\\ocr_engine.py 创建 PaddleOCR 的参数中加入：")
add_code(doc, "enable_mkldnn=False,")
add_text(doc, "RequestsDependencyWarning、没有 ccache、模型已缓存等通常不是致命错误。真正原因应查看 Traceback 的最后几行。")
add_heading(doc, "启动与健康检查", 2)
add_code(doc, 'cd /d "C:\\Users\\lenovo\\Desktop\\OCR total\\ocr-python-service"\nconda activate jrrg\npython app.py')
add_code(doc, "curl http://127.0.0.1:5001/health")
add_text(doc, "返回内容中的 ready 应为 true。")

add_heading(doc, "5. 启动 Java 后端")
add_code(doc, 'cd /d "C:\\Users\\lenovo\\Desktop\\OCR total\\ocr-business-server"\nmvn spring-boot:run')
add_text(doc, "出现 Tomcat started on port 8080 和 Started OcrBusinessApplication 即成功。终端不返回提示符是正常现象；Ctrl+C 用于停止。")

doc.add_section(WD_SECTION.NEW_PAGE)
add_heading(doc, "6. Node.js 与前端")
add_text(doc, "安装 Node.js LTS（Node.js 22 或更新的兼容 LTS），npm 会随 Node.js 一起安装：")
add_code(doc, "winget install --id OpenJS.NodeJS.LTS -e --source winget")
add_code(doc, "node --version\nnpm --version")
add_text(doc, "安装依赖并启动：")
add_code(doc, 'cd /d "C:\\Users\\lenovo\\Desktop\\OCR total\\ocr-web"\nnpm ci\nnpm run dev')
add_code(doc, "浏览器访问：http://localhost:5173")

add_heading(doc, "7. 每次运行项目的快捷流程")
add_heading(doc, "终端 1：Python OCR 服务", 2)
add_code(doc, 'conda activate jrrg\ncd /d "C:\\Users\\lenovo\\Desktop\\OCR total\\ocr-python-service"\npython app.py')
add_heading(doc, "终端 2：Java 后端", 2)
add_code(doc, 'cd /d "C:\\Users\\lenovo\\Desktop\\OCR total\\ocr-business-server"\nmvn spring-boot:run')
add_heading(doc, "终端 3：Vue 前端", 2)
add_code(doc, 'cd /d "C:\\Users\\lenovo\\Desktop\\OCR total\\ocr-web"\nnpm run dev')

add_heading(doc, "8. 常见故障")
add_bullets(doc, [
    "mvn 找不到：检查 MAVEN_HOME 和用户 Path，并彻底重启 VS Code。",
    "npm 找不到：安装 Node.js LTS，重启终端后检查 node 和 npm 版本。",
    "requirements.txt 找不到：进入 ocr-python-service，或使用完整路径。",
    "前端 HTTP 500：查看 Python 服务窗口的完整 Traceback；Java 通常只是转发错误。",
    "端口占用：检查 5001、8080、5173 是否已有旧服务运行。",
    "旧版 .doc 无法解析：安装 LibreOffice 并加入 Path；.docx 不需要。",
])
add_text(doc, "补充：项目默认使用本地 H2 数据库，无需安装 MySQL；上传文件默认限制 50 MB；首次 OCR 可能下载模型，需要联网且耗时较长。")

footer = doc.sections[-1].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(footer.add_run("OCR total 项目本地环境指南｜2026-08-03"), size=8, color=(100, 110, 120))

doc.save(OUTPUT)
print(OUTPUT)
