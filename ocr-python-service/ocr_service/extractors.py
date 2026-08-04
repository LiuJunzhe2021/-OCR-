from __future__ import annotations

import io
import shutil
import subprocess
import tempfile
from datetime import date, datetime
from pathlib import Path
from typing import Any

import fitz
import pandas as pd
import pdfplumber
import openpyxl
from docx import Document
from PIL import Image

from .config import Settings
from .ocr_engine import MultiModelOCREngine

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}


def format_cell(value: Any) -> str:
    if value is None or pd.isna(value):
        return ""
    if isinstance(value, pd.Timestamp):
        value = value.to_pydatetime()
    if isinstance(value, datetime):
        return value.isoformat(sep=" ", timespec="seconds")
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def section(
    source: str,
    method: str,
    text: str,
    metadata: dict[str, Any] | None = None,
    table_rows: list[list[str]] | None = None,
) -> dict[str, Any]:
    return {
        "source": source,
        "method": method,
        "text": text,
        "metadata": metadata or {},
        "tableRows": table_rows or [],
    }


class Extractors:
    def __init__(self, settings: Settings, engine: MultiModelOCREngine):
        self.settings = settings
        self.engine = engine

    def image(self, path: Path) -> tuple[list[dict[str, Any]], list[str]]:
        recognition = self.engine.recognize(path)
        return [
            section(
                "图片",
                f"ocr-{recognition.engine}",
                recognition.text,
                recognition.metadata(),
                [[line] for line in recognition.text.splitlines() if line.strip()],
            )
        ], recognition.warnings

    def excel(self, path: Path) -> tuple[list[dict[str, Any]], list[str]]:
        formula_cells: dict[str, list[str]] = {}
        if path.suffix.lower() == ".xlsx":
            workbook = openpyxl.load_workbook(path, data_only=False, read_only=True)
            try:
                for sheet in workbook.worksheets:
                    formula_cells[sheet.title] = [
                        cell.coordinate
                        for row in sheet.iter_rows()
                        for cell in row
                        if cell.data_type == "f"
                    ]
            finally:
                workbook.close()
        sheets = pd.read_excel(
            path,
            sheet_name=None,
            header=None,
            dtype=object,
            keep_default_na=False,
        )
        sections = []
        for name, frame in sheets.items():
            rows: list[list[str]] = []
            for values in frame.itertuples(index=False, name=None):
                row = [format_cell(value) for value in values]
                while row and row[-1] == "":
                    row.pop()
                if any(row):
                    rows.append(row)
            sections.append(
                section(
                    f"Excel工作表：{name}",
                    "pandas",
                    "\n".join("\t".join(row) for row in rows),
                    {
                        "sheet": name,
                        "rowCount": len(rows),
                        "columnCount": max((len(row) for row in rows), default=0),
                        "ocrUsed": False,
                        "formulaCells": formula_cells.get(name, []),
                    },
                    rows,
                )
            )
        return sections, []

    def pdf(self, path: Path, mode: str) -> tuple[list[dict[str, Any]], list[str]]:
        sections: list[dict[str, Any]] = []
        warnings: list[str] = []
        renderer = fitz.open(path)
        try:
            with pdfplumber.open(path) as document:
                for page_number, page in enumerate(document.pages, 1):
                    native_text = (
                        page.extract_text(x_tolerance=2, y_tolerance=3) or ""
                    ).strip()
                    meaningful = sum(
                        char.isalnum() or "\u4e00" <= char <= "\u9fff"
                        for char in native_text
                    )
                    scanned = mode == "ocr" or (
                        mode == "auto" and meaningful < self.settings.pdf_text_threshold
                    )
                    if mode == "native":
                        scanned = False

                    if not scanned:
                        sections.append(
                            section(
                                f"PDF第{page_number}页正文",
                                "pdfplumber-text",
                                native_text,
                                {
                                    "page": page_number,
                                    "pdfRoute": "native-page",
                                    "meaningfulCharacters": meaningful,
                                },
                                [[line] for line in native_text.splitlines() if line.strip()],
                            )
                        )
                        try:
                            for table_number, table in enumerate(page.extract_tables(), 1):
                                rows = [
                                    ["" if cell is None else str(cell).strip() for cell in row]
                                    for row in table
                                ]
                                rows = [row for row in rows if any(row)]
                                if rows:
                                    sections.append(
                                        section(
                                            f"PDF第{page_number}页表格{table_number}",
                                            "pdfplumber-table",
                                            "\n".join("\t".join(row) for row in rows),
                                            {
                                                "page": page_number,
                                                "table": table_number,
                                                "pdfRoute": "native-page",
                                            },
                                            rows,
                                        )
                                    )
                        except Exception as exc:
                            warnings.append(f"第{page_number}页表格提取失败：{exc}")
                        continue

                    dpi = min(max(self.settings.pdf_dpi, 150), 600)
                    scale = dpi / 72.0
                    pixmap = renderer[page_number - 1].get_pixmap(
                        matrix=fitz.Matrix(scale, scale), alpha=False
                    )
                    image = Image.open(io.BytesIO(pixmap.tobytes("png")))
                    try:
                        recognition = self.engine.recognize(image)
                    finally:
                        image.close()
                    metadata = {
                        "page": page_number,
                        "pdfRoute": "scanned-page",
                        "renderDpi": dpi,
                        "primaryModel": "paddleocr",
                        **recognition.metadata(),
                    }
                    rows = [
                        [value for value in line.split() if value]
                        for line in recognition.text.splitlines()
                        if line.strip()
                    ]
                    sections.append(
                        section(
                            f"PDF第{page_number}页扫描件",
                            f"ocr-{recognition.engine}",
                            recognition.text,
                            metadata,
                            rows,
                        )
                    )
                    warnings.extend(recognition.warnings)
        finally:
            renderer.close()
        return sections, warnings

    @staticmethod
    def _docx(path: Path) -> list[dict[str, Any]]:
        document = Document(path)
        paragraphs = [item.text for item in document.paragraphs if item.text.strip()]
        sections = [
            section(
                "Word正文",
                "python-docx",
                "\n".join(paragraphs),
                {},
                [[item] for item in paragraphs],
            )
        ]
        for number, table in enumerate(document.tables, 1):
            rows = [
                [cell.text.replace("\n", " ").strip() for cell in row.cells]
                for row in table.rows
            ]
            sections.append(
                section(
                    f"Word表格{number}",
                    "python-docx-table",
                    "\n".join("\t".join(row) for row in rows),
                    {"table": number},
                    rows,
                )
            )
        return sections

    def word(self, path: Path) -> tuple[list[dict[str, Any]], list[str]]:
        if path.suffix.lower() == ".docx":
            return self._docx(path), []
        executable = shutil.which("libreoffice") or shutil.which("soffice")
        if not executable:
            raise RuntimeError("解析.doc需要安装LibreOffice并加入PATH")
        with tempfile.TemporaryDirectory(prefix="doc_convert_") as temp_dir:
            process = subprocess.run(
                [
                    executable,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    temp_dir,
                    str(path),
                ],
                capture_output=True,
                text=True,
                timeout=120,
                check=False,
            )
            converted = Path(temp_dir) / f"{path.stem}.docx"
            if process.returncode != 0 or not converted.exists():
                raise RuntimeError(f"DOC转换失败：{process.stderr or process.stdout}")
            return self._docx(converted), ["旧版DOC已临时转换为DOCX"]
